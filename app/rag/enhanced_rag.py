"""Enhanced RAG Pipeline with Query Rewriting and Confidence Scoring"""

import os
import time
import asyncio
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import hashlib
from pathlib import Path

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.vectorstores import Chroma
from langchain.prompts import PromptTemplate

from app.llm.router import llm_router
from app.core.logging import logger
from app.core.cache import (
    get_cached_query_result,
    cache_query_result,
    get_cached_llm_response,
    cache_llm_response,
)
from app.core.async_utils import (
    ParallelSearchManager,
    AsyncBatchProcessor,
)
from app.core.adaptive_features import (
    get_optimal_search_config,
    response_time_tracker,
    adaptive_feature_manager,
)
from app.core.model_optimization import (
    get_optimized_embedding_model,
)
from app.core.monitoring import performance_monitor


class EnhancedRAGPipeline:
    """Enhanced RAG pipeline with query rewriting and confidence scoring"""

    def __init__(
        self,
        documents_dir: str = "data/documents",
        vectordb_dir: str = "data/vectordb",
        collection_name: str = "moji_documents",
        chunk_size: int = 800,  # Optimized for better chunking quality
        chunk_overlap: int = 150,  # Reduced overlap for efficiency
        use_faiss_fallback: bool = True,
        use_semantic_chunking: bool = True,
    ):
        self.documents_dir = Path(documents_dir)
        self.vectordb_dir = Path(vectordb_dir)
        self.collection_name = collection_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_faiss_fallback = use_faiss_fallback
        self.use_semantic_chunking = use_semantic_chunking
        self.faiss_retriever = None

        # Initialize components

        # Defer embedding model initialization until first use
        self.embeddings = None
        self._embeddings_initialized = False
        logger.info("Embedding model will be initialized on first use (lazy loading)")
        # Initialize text splitters with improved separators for Korean text
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n\n",   # Multiple line breaks (section boundaries)
                "\n\n",     # Paragraph boundaries
                "\n",       # Line breaks
                ". ",       # Sentence endings with space
                "。 ",      # Korean sentence endings
                "! ",       # Exclamations
                "? ",       # Questions
                "; ",       # Semicolons
                " ",        # Spaces
                "",         # Fallback
            ],
        )

        # Defer semantic chunker initialization until first use
        self.semantic_chunker = None
        self._semantic_chunker_initialized = False
        if self.use_semantic_chunking:
            logger.info(
                "Semantic chunker will be initialized on first use (lazy loading)"
            )

        # Defer vector store initialization until first use
        self.vectorstore = None
        self.faiss_retriever = None
        self._vectorstore_initialized = False
        self._is_persistent = False
        self._using_faiss = False
        logger.info("Vector store will be initialized on first use (lazy loading)")

        # 병렬 처리 관리자들
        self.parallel_search_manager = ParallelSearchManager(max_concurrent=4)
        self.batch_processor = AsyncBatchProcessor(batch_size=10, max_concurrent=3)

    def _initialize_semantic_chunker(self):
        """Initialize semantic chunker on first use (lazy loading)"""
        if self._semantic_chunker_initialized:
            return

        try:
            from app.rag.semantic_chunker import SemanticChunker, ChunkingStrategy

            self.semantic_chunker = SemanticChunker(
                min_chunk_size=max(200, self.chunk_size // 4),  # Balanced min size
                max_chunk_size=self.chunk_size,
                overlap_size=self.chunk_overlap,
                similarity_threshold=0.65,  # Higher threshold for better quality
                use_structure_hints=True,  # Enable structure-aware chunking
            )
            self.chunking_strategy = ChunkingStrategy.ADAPTIVE
            logger.info("Semantic chunker initialized successfully (lazy loading)")
            self._semantic_chunker_initialized = True
        except Exception as e:
            logger.warning(
                f"Failed to initialize semantic chunker: {e}, using standard chunking"
            )
            self.semantic_chunker = None
            self.use_semantic_chunking = False
            self._semantic_chunker_initialized = True

    def _initialize_embeddings(self):
        """Initialize embedding model on first use (lazy loading)"""
        if self._embeddings_initialized:
            return

        try:
            self.embeddings = get_optimized_embedding_model()
            logger.info("Using optimized embedding model (lazy loading)")
            self._embeddings_initialized = True
        except Exception as e:
            logger.warning(
                f"Failed to get optimized embedding model, using fallback: {e}"
            )
            from langchain_openai import OpenAIEmbeddings
            from app.core.config import settings

            api_key = settings.openai_api_key or settings.llm_api_key
            self.embeddings = OpenAIEmbeddings(
                api_key=api_key, model="text-embedding-3-small"
            )
            logger.info("Using basic OpenAI embedding model (lazy loading)")
            self._embeddings_initialized = True

    def _initialize_vectorstore(self):
        """Initialize vector store on first use (lazy loading)"""
        if self._vectorstore_initialized:
            return

        # Ensure embeddings are initialized first
        self._initialize_embeddings()

        try:
            # Try persistent storage first with proper directory setup
            persist_dir = Path(self.vectordb_dir)
            persist_dir.mkdir(parents=True, exist_ok=True)

            # Set proper permissions
            import os

            os.chmod(str(persist_dir), 0o755)

            # Initialize with persistent storage and telemetry disabled
            import chromadb
            from chromadb.config import Settings
            
            # Create ChromaDB client with telemetry disabled
            client_settings = Settings(
                anonymized_telemetry=False,
                allow_reset=True,
                is_persistent=True,
            )
            
            self.vectorstore = Chroma(
                collection_name=self.collection_name,
                embedding_function=self.embeddings,
                persist_directory=str(persist_dir),
                client_settings=client_settings,
            )

            logger.info(
                f"Using persistent Chroma store at: {persist_dir} (lazy loading)"
            )
            self._is_persistent = True
            self._using_faiss = False
            self._vectorstore_initialized = True

        except Exception as e:
            logger.warning(f"Failed to create persistent Chroma store: {e}")

            if self.use_faiss_fallback:
                logger.info("Attempting FAISS fallback")
                try:
                    from app.rag.retriever import VectorRetriever

                    self.faiss_retriever = VectorRetriever(
                        embeddings=self.embeddings,
                        index_path=str(self.vectordb_dir.parent / "faiss_index"),
                    )
                    logger.info(
                        "Successfully initialized FAISS fallback (lazy loading)"
                    )
                    self.vectorstore = None  # Use FAISS instead
                    self._is_persistent = True
                    self._using_faiss = True
                    self._vectorstore_initialized = True
                    return
                except Exception as faiss_error:
                    logger.warning(f"FAISS fallback also failed: {faiss_error}")

            logger.info("Falling back to in-memory Chroma store")
            # Final fallback to in-memory store
            self.vectorstore = Chroma(
                collection_name=self.collection_name,
                embedding_function=self.embeddings,
            )
            self._is_persistent = False
            self._using_faiss = False
            self._vectorstore_initialized = True

        # Query rewriting prompt
        self.query_rewrite_prompt = PromptTemplate(
            template="""당신은 한국어 검색 쿼리를 최적화하는 전문가입니다.
주어진 질문을 분석하여 관련 정보를 더 정확하게 찾을 수 있는 개선된 검색어를 생성하세요.

원본 질문: {query}

다음 전략을 적용하여 1개의 최적화된 검색어를 생성하세요:
1. 핵심 키워드 추출 및 확장 (동의어, 관련 용어 포함)
2. 불필요한 조사나 접속사 제거
3. 도메인 특화 용어 활용 (회사, 조직, 기술, 복지, 정책 등)
4. 구체적이고 명확한 표현 사용

개선된 검색어:""",
            input_variables=["query"],
        )

        # Answer generation prompt with confidence scoring
        self.answer_prompt = PromptTemplate(
            template="""당신은 주어진 컨텍스트를 바탕으로 정확한 답변을 제공하는 전문가입니다.
Chain-of-Thought 추론을 사용하여 체계적으로 답변을 생성하세요.

컨텍스트:
{context}

질문: {question}

다음 단계를 따라 답변하세요:

1. 분석: 질문의 핵심 요구사항을 파악하고 컨텍스트에서 관련 정보를 찾으세요.
2. 추론: 찾은 정보를 바탕으로 논리적으로 답변을 구성하세요.
3. 검증: 답변이 질문에 완전히 대답하는지, 컨텍스트와 일치하는지 확인하세요.

답변 형식:
답변: [컨텍스트 기반의 상세하고 정확한 답변]
신뢰도: [높음/중간/낮음]
근거: [답변의 근거가 된 주요 정보와 신뢰도 판단 이유]
출처: [사용된 문서 소스 목록]

컨텍스트에 관련 정보가 없다면 명확히 알려주세요.""",
            input_variables=["context", "question"],
        )

    async def load_documents(
        self, file_paths: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Load documents from files or directory"""
        try:
            documents = []
            processed_files = []

            if file_paths:
                # Load specific files
                for file_path in file_paths:
                    if os.path.exists(file_path):
                        docs = await self._load_single_document(file_path)
                        documents.extend(docs)
                        processed_files.append(file_path)
            else:
                # Load all documents from directory
                self.documents_dir.mkdir(parents=True, exist_ok=True)
                for file_path in self.documents_dir.glob("**/*"):
                    if file_path.is_file() and file_path.suffix in [
                        ".txt",
                        ".md",
                        ".docx",
                    ]:
                        docs = await self._load_single_document(str(file_path))
                        documents.extend(docs)
                        processed_files.append(str(file_path))

            # Add documents to vector store
            if documents:
                # Initialize vector store on first use
                self._initialize_vectorstore()
                # Filter complex metadata (None values, etc.) before storing

                filtered_documents = []

                for doc in documents:
                    # Clean metadata by removing None values and converting types
                    clean_metadata = {}
                    for key, value in doc.metadata.items():
                        if value is not None:
                            # Convert to acceptable types
                            if isinstance(value, (str, int, float, bool)):
                                clean_metadata[key] = value
                            else:
                                # Convert other types to string
                                clean_metadata[key] = str(value)

                    # Create new document with cleaned metadata
                    filtered_doc = Document(
                        page_content=doc.page_content, metadata=clean_metadata
                    )
                    filtered_documents.append(filtered_doc)

                if getattr(self, "_using_faiss", False) and self.faiss_retriever:
                    # Use FAISS
                    await self.faiss_retriever.create_index(filtered_documents)
                    logger.info(
                        f"Added {len(filtered_documents)} documents to FAISS index"
                    )
                else:
                    # Use ChromaDB
                    self.vectorstore.add_documents(filtered_documents)

                    # Note: ChromaDB 0.4.x+ automatically persists documents
                    if getattr(self, "_is_persistent", False):
                        logger.info(
                            f"Added {len(filtered_documents)} documents to persistent ChromaDB"
                        )
                    else:
                        logger.info(
                            f"Added {len(filtered_documents)} documents to in-memory ChromaDB"
                        )

            return {
                "success": True,
                "processed_files": processed_files,
                "total_chunks": len(documents),
                "message": f"Successfully processed {len(processed_files)} files into {len(documents)} chunks",
            }

        except Exception as e:
            logger.error(f"Error loading documents: {e}")
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to load documents",
            }

    async def _load_single_document(self, file_path: str) -> List[Document]:
        """Load and process a single document"""
        try:
            # Read file content
            if file_path.endswith(".docx"):
                # Use python-docx for Word documents
                try:
                    import docx

                    doc = docx.Document(file_path)
                    content = "\n".join(
                        [paragraph.text for paragraph in doc.paragraphs]
                    )
                except ImportError:
                    logger.warning("python-docx not installed, skipping .docx file")
                    return []
            else:
                # Text and markdown files
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()

            # Create document ID
            doc_id = hashlib.md5(file_path.encode()).hexdigest()[:8]

            # Choose chunking method
            if self.use_semantic_chunking:
                # Initialize semantic chunker on first use
                self._initialize_semantic_chunker()

            if self.use_semantic_chunking and self.semantic_chunker:
                # Use semantic chunking
                try:
                    documents = self.semantic_chunker.chunk_document(
                        content,
                        source_file=os.path.basename(file_path),
                        strategy=self.chunking_strategy,
                    )

                    # Add additional metadata
                    for i, doc in enumerate(documents):
                        doc.metadata.update(
                            {
                                "source": file_path,
                                "doc_id": doc_id,
                                "file_name": os.path.basename(file_path),
                                "created_at": datetime.utcnow().isoformat(),
                            }
                        )
                        # Convert ChunkingStrategy to string if present
                        if "chunk_type" in doc.metadata and hasattr(
                            doc.metadata["chunk_type"], "value"
                        ):
                            doc.metadata["chunk_type"] = doc.metadata[
                                "chunk_type"
                            ].value

                    # Log chunking stats
                    if documents:
                        stats = self.semantic_chunker.get_chunking_stats(documents)
                        logger.info(f"Semantic chunking stats for {file_path}: {stats}")

                except Exception as e:
                    logger.warning(
                        f"Semantic chunking failed for {file_path}: {e}, using fallback"
                    )
                    # Fallback to standard chunking
                    chunks = self.text_splitter.split_text(content)
                    documents = self._create_standard_documents(
                        chunks, file_path, doc_id
                    )
            else:
                # Use standard chunking
                chunks = self.text_splitter.split_text(content)
                documents = self._create_standard_documents(chunks, file_path, doc_id)

            logger.info(f"Loaded {len(documents)} chunks from {file_path}")
            return documents

        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return []

    def _create_standard_documents(
        self, chunks: List[str], file_path: str, doc_id: str
    ) -> List[Document]:
        """Create Document objects from text chunks using standard method"""
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "doc_id": doc_id,
                    "chunk_id": f"{doc_id}_{i}",
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_name": os.path.basename(file_path),
                    "created_at": datetime.utcnow().isoformat(),
                    "chunk_type": "standard",
                },
            )
            documents.append(doc)
        return documents

    async def rewrite_query(self, query: str, force_rewrite: bool = False) -> List[str]:
        """Rewrite query to improve search results"""
        try:
            # 적응형 결정: 쿼리 재작성 필요 여부 확인
            if not force_rewrite:
                optimal_config = get_optimal_search_config(query)
                if not optimal_config["use_query_rewriting"]:
                    logger.info("Query rewriting skipped based on adaptive analysis")
                    return [query]

            # Get LLM
            llm = await llm_router.get_langchain_model()

            # Create chain using RunnableSequence (new LangChain pattern)
            chain = self.query_rewrite_prompt | llm

            # Generate alternative queries
            result = await chain.ainvoke({"query": query})

            # Parse results (new LangChain format returns content directly)
            result_text = result.content if hasattr(result, "content") else str(result)
            alternative_queries = [
                q.strip() for q in result_text.strip().split("\n") if q.strip()
            ]

            # Add original query
            all_queries = [query] + alternative_queries
            
            # Add keyword-based variations
            all_queries.extend(self._generate_keyword_variations(query))
            
            # Remove duplicates while preserving order
            seen = set()
            unique_queries = []
            for q in all_queries:
                if q not in seen:
                    seen.add(q)
                    unique_queries.append(q)

            logger.info(f"Generated {len(unique_queries)} query variations")
            return unique_queries[:5]  # Limit to top 5 variations

        except Exception as e:
            logger.error(f"Error rewriting query: {e}")
            return [query]  # Return original query on error
    
    def _generate_keyword_variations(self, query: str) -> List[str]:
        """Generate keyword-based query variations"""
        variations = []
        
        # Domain-specific synonym mappings
        synonyms = {
            "회사": ["기업", "조직", "MOJI"],
            "복지": ["복리후생", "혜택", "베네핏"],
            "휴가": ["연차", "휴일", "쉬는날"],
            "기술": ["기술스택", "스택", "사용기술"],
            "개발": ["코딩", "프로그래밍", "구현"],
            "배포": ["디플로이", "deployment", "릴리즈"],
            "점심": ["점심값", "식대", "식사"],
        }
        
        # Check for keywords and add variations
        for keyword, syns in synonyms.items():
            if keyword in query:
                for syn in syns:
                    variations.append(query.replace(keyword, syn))
        
        return variations

    async def search_with_rewriting(
        self,
        query: str,
        # 검색 및 작성에 최적화를 위한 Threshold 값 조정 테스트
        k: int = 5,  # 문서 참조 걔수
        score_threshold: float = 2.0,  # 한국어 지원을 위해 임계값 상향 조정
        use_parallel_search: bool = True,
    ) -> Tuple[List[Document], Dict[str, Any]]:
        """Search with query rewriting and relevance scoring"""
        try:
            # Rewrite query
            queries = await self.rewrite_query(query)

            # Search with all query variations (병렬 또는 순차)
            if use_parallel_search and len(queries) > 1:
                all_results = await self._parallel_vector_search(
                    queries, k, score_threshold
                )
            else:
                all_results = await self._sequential_vector_search(
                    queries, k, score_threshold
                )

            # Sort by score (ascending - lower distance is better)
            all_results.sort(key=lambda x: x[1])

            # Take top k results
            top_results = all_results[:k]

            # Extract documents and metadata
            documents = [doc for doc, _, _ in top_results]
            search_metadata = {
                "original_query": query,
                "rewritten_queries": queries,
                "total_results": len(all_results),
                "returned_results": len(documents),
                "score_threshold": score_threshold,
                "result_details": [
                    {
                        "chunk_id": doc.metadata.get("chunk_id"),
                        "source": doc.metadata.get("source"),
                        "score": score,
                        "matched_query": matched_query,
                    }
                    for doc, score, matched_query in top_results
                ],
            }

            return documents, search_metadata

        except Exception as e:
            logger.error(f"Error in search with rewriting: {e}")
            return [], {"error": str(e)}

    async def _parallel_vector_search(
        self, queries: List[str], k: int, score_threshold: float
    ) -> List[Tuple[Document, float, str]]:
        """병렬 벡터 검색 실행"""
        try:
            # 각 쿼리에 대한 검색 태스크 생성
            search_tasks = []
            for query in queries:
                task = self._single_vector_search(query, k * 2, score_threshold)
                search_tasks.append(task)

            # 병렬 실행 (타임아웃 추가)
            start_time = time.time()
            try:
                results_per_query = await asyncio.wait_for(
                    asyncio.gather(*search_tasks, return_exceptions=True),
                    timeout=10.0,  # 10초 타임아웃
                )
            except asyncio.TimeoutError:
                logger.warning("Parallel search timeout, falling back to sequential")
                return await self._sequential_vector_search(queries, k, score_threshold)
            elapsed_time = time.time() - start_time

            logger.info(
                f"Parallel vector search completed in {elapsed_time:.3f}s for {len(queries)} queries"
            )

            # 결과 통합 및 중복 제거
            all_results = []
            seen_chunks = set()

            for i, results in enumerate(results_per_query):
                if isinstance(results, Exception):
                    logger.error(f"Search failed for query '{queries[i]}': {results}")
                    continue

                for doc, score in results:
                    chunk_id = doc.metadata.get("chunk_id")
                    if chunk_id not in seen_chunks:
                        seen_chunks.add(chunk_id)
                        all_results.append((doc, score, queries[i]))

            return all_results

        except Exception as e:
            logger.error(f"Parallel vector search error: {e}")
            # 폴백: 순차 검색
            return await self._sequential_vector_search(queries, k, score_threshold)

    async def _sequential_vector_search(
        self, queries: List[str], k: int, score_threshold: float
    ) -> List[Tuple[Document, float, str]]:
        """순차 벡터 검색 실행"""
        all_results = []
        seen_chunks = set()

        for q in queries:
            try:
                results = await self._single_vector_search(q, k * 2, score_threshold)

                for doc, score in results:
                    chunk_id = doc.metadata.get("chunk_id")
                    if chunk_id not in seen_chunks:
                        seen_chunks.add(chunk_id)
                        all_results.append((doc, score, q))

            except Exception as e:
                logger.error(f"Sequential search failed for query '{q}': {e}")
                continue

        return all_results

    async def _single_vector_search(
        self, query: str, k: int, score_threshold: float
    ) -> List[Tuple[Document, float]]:
        """단일 쿼리 벡터 검색"""
        try:
            # Initialize vector store on first use
            self._initialize_vectorstore()
            # Check if using FAISS or ChromaDB
            if getattr(self, "_using_faiss", False) and self.faiss_retriever:
                # Use FAISS retriever - call the underlying sync method directly
                try:
                    # Access the underlying sync search method directly
                    if (
                        hasattr(self.faiss_retriever, "vector_store")
                        and self.faiss_retriever.vector_store
                    ):
                        loop = asyncio.get_event_loop()
                        faiss_results = await loop.run_in_executor(
                            None,
                            lambda: self.faiss_retriever.vector_store.similarity_search_with_score(
                                query, k=k
                            ),
                        )
                        # Convert FAISS results to (doc, score) tuples
                        if faiss_results:
                            results = [
                                (doc, 1.0 - score) for doc, score in faiss_results
                            ]  # Convert similarity to distance
                        else:
                            results = []
                    else:
                        logger.warning("FAISS vector store not available")
                        results = []
                except Exception as faiss_error:
                    logger.warning(
                        f"FAISS search failed: {faiss_error}, falling back to empty results"
                    )
                    results = []
            else:
                # Use ChromaDB with enhanced error handling and debugging
                try:
                    # First check if vectorstore is properly initialized
                    if not self.vectorstore:
                        logger.error("Vectorstore is not initialized")
                        return []

                    # Check if collection has any documents
                    try:
                        collection_count = self.vectorstore._collection.count()
                        logger.debug(f"Collection count: {collection_count}")
                        if collection_count == 0:
                            logger.warning("Vector collection is empty")
                            return []
                    except Exception as count_error:
                        logger.warning(
                            f"Could not check collection count: {count_error}"
                        )

                    # Debug logging for vectorstore state
                    logger.debug(f"Vectorstore type: {type(self.vectorstore)}")
                    logger.debug(
                        f"Collection name: {getattr(self.vectorstore, '_collection_name', 'Unknown')}"
                    )

                    # Perform search with detailed error tracking
                    logger.debug(
                        f"Starting similarity search for query: '{query[:50]}...'"
                    )
                    loop = asyncio.get_event_loop()

                    try:
                        results = await loop.run_in_executor(
                            None,
                            lambda: self.vectorstore.similarity_search_with_score(
                                query, k=k
                            ),
                        )
                        logger.debug(
                            f"Search completed, result type: {type(results)}, length: {len(results) if results else 0}"
                        )
                    except AttributeError as attr_error:
                        logger.error(
                            f"AttributeError in similarity search - likely a .get() call on None: {attr_error}"
                        )
                        logger.error(f"Query: {query}")
                        logger.error(
                            f"Vectorstore state: {vars(self.vectorstore) if hasattr(self.vectorstore, '__dict__') else 'No __dict__'}"
                        )
                        # Try to identify the specific None object
                        if "'NoneType' object has no attribute 'get'" in str(
                            attr_error
                        ):
                            logger.error(
                                "IDENTIFIED: This is the NoneType .get() error we're tracking!"
                            )
                            logger.error(
                                "Likely causes: cache system, metadata handling, or internal ChromaDB state"
                            )
                        raise attr_error
                    except Exception as search_error:
                        logger.error(
                            f"Other search error: {type(search_error).__name__}: {search_error}"
                        )
                        raise search_error

                    # Ensure results is a list
                    if not results:
                        logger.debug("Search returned empty results")
                        results = []
                    elif not isinstance(results, list):
                        logger.error(f"Unexpected search result type: {type(results)}")
                        results = []
                    else:
                        logger.debug(f"Search returned {len(results)} results")

                except Exception as search_error:
                    logger.error(
                        f"ChromaDB search failed: {type(search_error).__name__}: {search_error}"
                    )
                    # Add stack trace for debugging
                    import traceback

                    logger.error(f"Full traceback:\n{traceback.format_exc()}")
                    return []

            # 점수 필터링
            filtered_results = []
            if results:
                for item in results:
                    try:
                        if isinstance(item, tuple) and len(item) == 2:
                            doc, score = item
                            if score <= score_threshold:
                                filtered_results.append((doc, score))
                    except Exception as filter_error:
                        logger.warning(f"Error filtering result: {filter_error}")
                        continue

            return filtered_results

        except Exception as e:
            logger.error(f"Single vector search error for query '{query}': {e}")
            import traceback

            traceback.print_exc()
            return []

    async def answer_with_confidence(
        self,
        query: str,
        # 답변 신뢰도도 최적화를 위한 Threshold 값 조정 테스트
        k: int = 5,
        score_threshold: float = 1.5,  # 한국어 지원을 위해 임계값 상향 조정
        context: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Generate answer with confidence scoring and source citations"""
        # 성능 추적 시작
        request_id = f"rag_{int(time.time() * 1000)}"
        response_time_tracker.start_tracking(request_id)
        start_time = time.time()

        # 모니터링 시작
        monitoring_request_id = performance_monitor.record_request_start(
            request_id, "rag_query"
        )

        try:
            # 1. 쿼리 분석 및 최적 설정 결정
            optimal_config = get_optimal_search_config(query, context)
            logger.info(f"Adaptive config: {optimal_config}")

            # 2. 캐시 조회 먼저 시도
            search_params = {
                "k": k,
                "score_threshold": score_threshold,
                "config": optimal_config,
            }
            cached_result = await get_cached_query_result(query, search_params)
            if cached_result:
                # 캐시 히트 시간 기록
                elapsed_time = response_time_tracker.end_tracking(request_id)
                adaptive_feature_manager.record_performance(elapsed_time)
                logger.info(
                    f"Cache hit for query: {query[:50]}... ({elapsed_time:.3f}s)"
                )
                return cached_result

            # 3. 적응형 설정에 따른 문서 검색
            documents, search_metadata = await self.search_with_rewriting(
                query,
                k=k,
                score_threshold=score_threshold,
                use_parallel_search=optimal_config["use_parallel_search"],
            )

            # 검색 메타데이터에 적응형 정보 추가
            search_metadata["adaptive_config"] = optimal_config

            if not documents:
                return {
                    "answer": "죄송합니다. 업로드된 문서에서 관련 정보를 찾을 수 없습니다. 다른 질문을 해주시거나 관련 문서를 업로드해 주세요.",
                    "confidence": "LOW",
                    "reasoning": "문서 기반 검색에서 관련 정보를 찾지 못했습니다",
                    "sources": [],
                    "search_metadata": search_metadata,
                }

            # Prepare context
            context_parts = []
            for i, doc in enumerate(documents):
                source = os.path.basename(doc.metadata.get("source", "Unknown"))
                context_parts.append(f"[Source {i+1}: {source}]\n{doc.page_content}")

            context = "\n\n".join(context_parts)

            # LLM 응답 캐시 조회
            context_len = len(context) if context else 0
            llm_params = {"model": "default", "context_length": context_len}
            context_prefix = context[:500] if context else ""
            cached_llm_response = await get_cached_llm_response(
                f"{query}||{context_prefix}", llm_params
            )

            if cached_llm_response:
                answer_text = str(cached_llm_response)
                logger.info(f"LLM cache hit for query: {query[:30]}...")
            else:
                # Generate answer with simpler prompt
                try:
                    llm = await llm_router.get_langchain_model()

                    # Use a strict prompt for document-based answers only
                    simple_prompt = PromptTemplate(
                        template="""다음 문맥에 명시된 정보만을 바탕으로 질문에 정확하게 답변해주세요.

중요한 지침:
- 문맥에 명시된 정보만 사용하세요
- 문맥에 없는 내용은 절대 추측하거나 추가하지 마세요
- 관련 정보가 부족하면 "제공된 문서에서 해당 정보를 찾을 수 없습니다"라고 답변하세요
- 일반적인 지식이나 추정을 사용하지 마세요

문맥:
{context}

질문: {question}

답변:""",
                        input_variables=["context", "question"],
                    )

                    chain = simple_prompt | llm
                    result = await chain.ainvoke(
                        {"context": context, "question": query}
                    )

                    # Check if result is valid (new LangChain format)
                    if not result:
                        logger.error(f"Invalid LLM result: {result}")
                        answer_text = (
                            "죄송합니다. 답변을 생성하는 중 오류가 발생했습니다."
                        )
                    else:
                        # Get the answer text (new format)
                        if hasattr(result, "content"):
                            answer_text = str(result.content).strip()
                        else:
                            answer_text = str(result).strip()

                        # LLM 응답 캐시 저장 (1시간 TTL)
                        await cache_llm_response(
                            f"{query}||{context_prefix}",
                            answer_text,
                            llm_params,
                            ttl=3600,
                        )

                except Exception as llm_error:
                    logger.error(f"LLM generation failed: {llm_error}")
                    answer_text = f"답변 생성 중 오류가 발생했습니다: {str(llm_error)}"

            # Simple confidence scoring based on answer length and content
            confidence = (
                "HIGH"
                if len(answer_text) > 100
                else "MEDIUM" if len(answer_text) > 50 else "LOW"
            )

            # Map sources to actual file paths
            source_files = list(
                set([doc.metadata.get("source", "Unknown") for doc in documents])
            )

            # 최종 결과 구성
            final_result = {
                "answer": answer_text,
                "confidence": confidence,
                "reasoning": f"답변이 {len(documents)}개의 관련 문서를 바탕으로 생성되었습니다.",
                "sources": source_files,
                "search_metadata": search_metadata,
                "context_used": len(documents),
                "timestamp": datetime.utcnow().isoformat(),
                "processing_time": f"{time.time() - start_time:.3f}s",
            }

            # 성능 추적 완료
            elapsed_time = response_time_tracker.end_tracking(request_id)
            adaptive_feature_manager.record_performance(elapsed_time)

            # 모니터링 완료 (성공)
            performance_monitor.record_request_end(
                monitoring_request_id,
                success=True,
                operation="rag_query",
                additional_metrics={
                    "documents_retrieved": len(documents),
                    "processing_mode": optimal_config["complexity"],
                    "cache_used": "cache_hit" in final_result.get("reasoning", ""),
                },
            )

            # 최종 결과에 성능 정보 추가
            final_result.update(
                {
                    "adaptive_config": optimal_config,
                    "actual_processing_time": f"{elapsed_time:.3f}s",
                    "estimated_vs_actual": f"estimated: {optimal_config['estimated_time']:.1f}s, actual: {elapsed_time:.1f}s",
                    "performance_mode": optimal_config["processing_mode"],
                }
            )

            # 결과를 캐시에 저장 (30분 TTL)
            await cache_query_result(query, final_result, search_params, ttl=1800)
            logger.info(
                f"Adaptive query processed: {query[:50]}... ({elapsed_time:.3f}s, mode: {optimal_config['processing_mode']})"
            )

            return final_result

        except Exception as e:
            logger.error(f"Error generating answer: {e}")

            # 모니터링 완료 (실패)
            performance_monitor.record_request_end(
                monitoring_request_id, success=False, operation="rag_query"
            )

            return {
                "answer": f"오류가 발생했습니다: {str(e)}",
                "confidence": "LOW",
                "reasoning": "Error occurred during processing",
                "sources": [],
                "error": str(e),
            }

    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector store collection"""
        try:
            # Initialize vector store on first use
            self._initialize_vectorstore()
            # Get collection
            collection = self.vectorstore._collection

            # Get stats
            count = collection.count()

            return {
                "collection_name": self.collection_name,
                "total_documents": count,
                "vectordb_path": str(self.vectordb_dir),
                "embedding_model": "text-embedding-3-small",
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
            }

        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {"error": str(e)}


# Global instance
rag_pipeline = EnhancedRAGPipeline()

# Hybrid search integration
_hybrid_pipeline = None


def get_hybrid_pipeline():
    """하이브리드 검색 파이프라인 인스턴스 반환"""
    global _hybrid_pipeline
    if _hybrid_pipeline is None:
        from app.rag.hybrid_search import HybridRAGPipeline

        _hybrid_pipeline = HybridRAGPipeline(rag_pipeline)
    return _hybrid_pipeline
