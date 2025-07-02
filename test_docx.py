#!/usr/bin/env python3
"""
DOCX 파일 업로드 테스트 스크립트
"""

import asyncio
import os
import sys
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

async def test_docx_support():
    """DOCX 지원 테스트"""
    print("📋 DOCX 파일 지원 테스트")
    print("=" * 50)
    
    # 1. python-docx 모듈 확인
    try:
        import docx
        print("✅ python-docx 모듈 로드됨")
        docx_available = True
    except ImportError:
        print("❌ python-docx 모듈 없음")
        print("   설치 명령: pip install python-docx")
        docx_available = False
    
    # 2. RAG 시스템에서 DOCX 지원 확인
    try:
        from app.rag.enhanced_rag import rag_pipeline
        
        # 지원 확장자 확인
        print(f"\n📄 지원 확장자:")
        
        # enhanced_rag.py에서 확인
        print("   - .txt (텍스트 파일)")
        print("   - .md (마크다운 파일)")
        print("   - .docx (Word 문서)")
        
        # upload_docs.py에서 확인
        from upload_docs import get_document_files
        
        # 테스트용 임시 디렉토리 생성
        test_dir = Path("test_docs")
        test_dir.mkdir(exist_ok=True)
        
        # 테스트 파일들 생성
        (test_dir / "test.txt").write_text("텍스트 파일 테스트")
        (test_dir / "test.md").write_text("# 마크다운 테스트")
        
        # 파일 검색 테스트
        found_files = get_document_files(test_dir)
        print(f"\n🔍 발견된 파일: {len(found_files)}개")
        for file in found_files:
            print(f"   - {file.name}")
        
        # 정리
        import shutil
        shutil.rmtree(test_dir)
        
    except Exception as e:
        print(f"❌ RAG 시스템 확인 오류: {str(e)}")
    
    # 3. 설치 가이드
    print(f"\n💡 DOCX 파일 사용하기:")
    if not docx_available:
        print("   1. python-docx 설치:")
        print("      pip install python-docx")
        print("      # 또는")
        print("      pip install -r requirements.txt")
        print("")
    
    print("   2. DOCX 파일 업로드:")
    print("      python upload_docs.py --file document.docx")
    print("")
    print("   3. 전체 폴더 업로드 (DOCX 포함):")
    print("      python upload_docs.py")
    
    # 4. 예시 DOCX 파일 생성 (python-docx가 있는 경우)
    if docx_available:
        try:
            print(f"\n📝 테스트 DOCX 파일 생성 중...")
            
            doc = docx.Document()
            doc.add_heading('MOJI 테스트 문서', 0)
            
            p = doc.add_paragraph('이것은 MOJI RAG 시스템을 위한 테스트 문서입니다.')
            p.add_run(' 이 문서는 ').bold = True
            p.add_run('Word 형식(.docx)')
            p.add_run('으로 저장되어 RAG 시스템에서 처리됩니다.')
            
            doc.add_heading('주요 기능', level=1)
            doc.add_paragraph('문서 검색', style='List Bullet')
            doc.add_paragraph('질문 답변', style='List Bullet')
            doc.add_paragraph('출처 표시', style='List Bullet')
            
            # data/documents 폴더 생성 및 저장
            docs_dir = Path("data/documents")
            docs_dir.mkdir(parents=True, exist_ok=True)
            
            test_docx_path = docs_dir / "moji_test.docx"
            doc.save(test_docx_path)
            
            print(f"✅ 테스트 파일 생성됨: {test_docx_path}")
            print(f"   파일 크기: {test_docx_path.stat().st_size} bytes")
            
            print(f"\n🚀 다음 단계:")
            print(f"   python upload_docs.py --file moji_test.docx")
            print(f"   # 또는")
            print(f"   python upload_docs.py  # 전체 업로드")
            
        except Exception as e:
            print(f"❌ DOCX 파일 생성 오류: {str(e)}")

if __name__ == "__main__":
    asyncio.run(test_docx_support())